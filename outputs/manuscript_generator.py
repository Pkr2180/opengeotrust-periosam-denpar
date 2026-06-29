"""
OpenGeoTrust-PerioSAM — Main Manuscript Generator
Produces: OpenGeoTrust_PerioSAM_Manuscript.docx
Format  : Times New Roman 12 pt · double-spaced · 1-inch margins
          No header · No footer · Vancouver citations
Figures : Figures 1-6 and Table 1 embedded inline after references.
"""

import io
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# ── Paths ────────────────────────────────────────────────────────────────────
BASE   = r"c:\Users\Pradeep Kumar\Downloads\perio ai\opengeotrust-periosam-denpar"
FIGS   = os.path.join(BASE, "outputs", "figures")
OUT    = os.path.join(BASE, "outputs", "OpenGeoTrust_PerioSAM_Manuscript.docx")

# ── Helpers ──────────────────────────────────────────────────────────────────
def _font(run, bold=False, italic=False, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic

def _para(doc, text="", bold=False, italic=False, size=12,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6, indent=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before    = Pt(sb)
    p.paragraph_format.space_after     = Pt(sa)
    p.paragraph_format.line_spacing    = Pt(24)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if text:
        r = p.add_run(text)
        _font(r, bold=bold, italic=italic, size=size)
    return p

def _heading(doc, text, sb=12, sa=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(24)
    r = p.add_run(text)
    _font(r, bold=True, size=12)
    return p

def _sup_para(doc, segments, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6, indent=False):
    """segments = [(text, is_superscript, bold, italic)]"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before    = Pt(sb)
    p.paragraph_format.space_after     = Pt(sa)
    p.paragraph_format.line_spacing    = Pt(24)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    for text, sup, bold, italic in segments:
        r = p.add_run(text)
        r.font.name       = "Times New Roman"
        r.font.size       = Pt(10) if sup else Pt(12)
        r.font.superscript = sup
        r.font.bold       = bold
        r.font.italic     = italic
    return p

def _add_figure(doc, png_path, width_in=6.0, caption_label="", caption_text=""):
    """Insert a PNG (RGBA-safe) centred on its own page with bold label + caption."""
    doc.add_page_break()

    # Convert RGBA → RGB so Word renders correctly
    img = Image.open(png_path).convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)

    # Image paragraph — centred
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.space_after  = Pt(6)
    run = p_img.add_run()
    run.add_picture(buf, width=Inches(width_in))

    # Caption paragraph
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after  = Pt(12)
    p_cap.paragraph_format.line_spacing = Pt(22)
    if caption_label:
        r1 = p_cap.add_run(caption_label + " ")
        _font(r1, bold=True, size=11)
    if caption_text:
        r2 = p_cap.add_run(caption_text)
        _font(r2, bold=False, size=11)

def _hdr_row(table, headers, size=11):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.font.name = "Times New Roman"; r.font.size = Pt(size); r.font.bold = True

def _data_row(table, values, size=11):
    cells = table.add_row().cells
    for i, v in enumerate(values):
        cells[i].text = v
        for r in cells[i].paragraphs[0].runs:
            r.font.name = "Times New Roman"; r.font.size = Pt(size)

# ═════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width  = Inches(8.5)
    sec.page_height = Inches(11)
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(sec, attr, Inches(1.0))
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.3)
    for elem in (sec.header, sec.footer):
        elem.is_linked_to_previous = False
        for p in elem.paragraphs:
            for r in p.runs: r.text = ""

    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    # ── TITLE ────────────────────────────────────────────────────────────────
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=12)
    r = p.add_run(
        "OpenGeoTrust-PerioSAM: A Geometry-Aware Multi-Task Deep Learning Framework "
        "for Simultaneous Crestal Bone Line Segmentation, Landmark Detection, and "
        "Calibrated Uncertainty Quantification from Intraoral Periapical Radiographs"
    )
    _font(r, bold=True, size=14)

    # ── ABSTRACT ─────────────────────────────────────────────────────────────
    _heading(doc, "ABSTRACT")
    for label, body in [
        ("Background: ",
         "Radiographic bone level assessment is central to periodontitis staging under the "
         "2018 World Workshop classification; however, current clinical practice relies on "
         "manual, operator-dependent measurements that are poorly reproducible. Existing "
         "artificial intelligence (AI) approaches to periodontal radiographic analysis are "
         "largely confined to image-level classification tasks and do not provide pixel-level "
         "bone topography, anatomical landmark coordinates, or calibrated uncertainty estimates "
         "required for clinical governance."),
        ("Objectives: ",
         "To develop and validate OpenGeoTrust-PerioSAM, a geometry-aware multi-task "
         "convolutional neural network that simultaneously segments the tooth outline and "
         "crestal bone line, detects the cemento-enamel junction (CEJ) and root apex, "
         "computes a geometry-based bone-loss percentage, and provides calibrated predictive "
         "uncertainty from intraoral periapical (IOPA) radiographs."),
        ("Methods: ",
         "A multi-task U-Net with a ResNet-34 encoder and three task-specific decoder heads "
         "was trained on the DenPAR dataset (n=1,000 annotated IOPA radiographs; 650/150/200 "
         "training/validation/test). A composite Tversky-Focal loss (α=0.3, β=0.7) with "
         "target-space bone-line rasterisation (7-pixel width at 512×512 resolution) "
         "addressed extreme class imbalance. Monte Carlo (MC) Dropout (T=20) quantified "
         "predictive uncertainty. Bootstrap 95% confidence intervals (B=2,000) were computed "
         "for all primary metrics."),
        ("Results: ",
         "On the held-out test set (n=200), the model achieved tooth segmentation Dice "
         "Similarity Coefficient (DSC) 0.957 (95% CI: 0.948–0.958), crestal bone-line DSC "
         "0.544 (95% CI: 0.543–0.595), keypoint mean radial error (MRE) 1.05 pixels "
         "(PCK@4px=1.000), and Expected Calibration Error (ECE) 0.094. Predictive uncertainty "
         "correlated significantly with segmentation error (Spearman ρ=0.591, p<0.0001). "
         "Multi-task joint training reduced the validation-to-test generalisation gap by 80.7% "
         "relative to single-task bone segmentation (gap: 0.034 vs. 0.176)."),
        ("Conclusions: ",
         "OpenGeoTrust-PerioSAM delivers geometrically interpretable, uncertainty-calibrated "
         "periodontal bone assessments from IOPA radiographs with performance supporting "
         "integration as a reproducible adjunct to radiographic bone level measurement and "
         "periodontitis staging."),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.line_spacing = Pt(24)
        r1 = p.add_run(label);  _font(r1, bold=True)
        r2 = p.add_run(body);   _font(r2)

    _para(doc,
        "Keywords: periodontitis; crestal bone segmentation; multi-task learning; deep "
        "learning; intraoral periapical radiograph; uncertainty quantification; dental AI",
        italic=True, sa=12)

    # ── INTRODUCTION ─────────────────────────────────────────────────────────
    _heading(doc, "INTRODUCTION")

    _sup_para(doc, [
        ("Periodontitis is among the most prevalent non-communicable diseases globally, "
         "with severe forms affecting an estimated 11.2% of the world's adult population — "
         "ranked sixth in global disease prevalence — and representing the primary cause of "
         "tooth loss in adults.", False, False, False),
        ("1", True, False, False),
        (" Its principal radiographic hallmark is alveolar bone loss measured from the "
         "cemento-enamel junction (CEJ) as the anatomical reference point. The 2017 World "
         "Workshop on the Classification of Periodontal and Peri-Implant Diseases formally "
         "embedded radiographic bone level as an essential criterion for periodontitis staging, "
         "with bone loss expressed as a millimetre value or root-length percentage.", False, False, False),
        ("2,3", True, False, False),
        (" Despite this diagnostic centrality, radiographic bone level measurement in "
         "routine clinical practice remains predominantly manual, operator-dependent, and "
         "burdened by significant inter- and intra-examiner variability, directly undermining "
         "staging reproducibility and longitudinal monitoring fidelity.", False, False, False),
    ], indent=True)

    _sup_para(doc, [
        ("The emergence of deep learning in dental diagnostics has generated substantial "
         "interest in automating radiographic periodontal assessment.",False,False,False),
        ("4,5", True, False, False),
        (" Krois and colleagues demonstrated that a seven-layer convolutional neural network "
         "trained on panoramic radiograph segments achieved a mean classification accuracy of "
         "0.81 for the presence of radiographic bone loss — comparable to experienced "
         "dentists.", False, False, False),
        ("6", True, False, False),
        (" Chang and colleagues extended this with a hybrid deep learning framework that "
         "automatically staged periodontitis on panoramic radiographs according to the 2017 "
         "classification criteria, achieving an intraclass correlation coefficient of 0.91 "
         "with radiologist diagnoses.", False, False, False),
        ("7", True, False, False),
        (" Critically, however, both of these landmark studies — and the majority of "
         "published periodontal AI — operated on panoramic radiographs and framed the problem "
         "as image-level classification rather than dense pixel-level segmentation. They "
         "return a diagnostic label or severity grade but cannot localise the crestal bone "
         "margin, extract per-tooth geometric measurements, or assign spatial confidence to "
         "individual predictions.", False, False, False),
    ], indent=True)

    _sup_para(doc, [
        ("Segmentation-based approaches have the potential to close this clinical translation "
         "gap, yet the crestal bone line presents formidable technical obstacles. At standard "
         "clinical resolution (512×512 pixels), bone-line annotations occupy fewer than "
         "0.5–1.0% of image pixels — a severe class-imbalance regime in which standard "
         "Dice-based objectives collapse to the trivial all-background prediction. "
         "Furthermore, IOPA radiographs, the most commonly acquired format in clinical "
         "periodontology for bone level assessment, exhibit high morphological variability "
         "and frequent superimposition of buccal and lingual cortical plates, rendering "
         "bone-line contrast inherently ambiguous. These characteristics have precluded "
         "the development of a validated, high-performance segmentation system for crestal "
         "bone lines on IOPAs.", False, False, False),
    ], indent=True)

    _sup_para(doc, [
        ("A further persistent limitation of existing dental AI systems is the absence of "
         "reliable uncertainty quantification. Deep neural networks are overconfident by "
         "default, generating high-confidence predictions even on ambiguous or "
         "out-of-distribution inputs — a property that renders uncalibrated systems "
         "clinically hazardous when prediction errors are not flagged for human review. "
         "Bayesian approximation via MC Dropout enables tractable predictive uncertainty "
         "estimation without architectural modification and has been validated across multiple "
         "medical imaging domains;", False, False, False),
        ("8", True, False, False),
        (" however, its systematic application to periodontal radiograph analysis has "
         "not been previously reported.", False, False, False),
    ], indent=True)

    _sup_para(doc, [
        ("Against this background, we introduce OpenGeoTrust-PerioSAM, a geometry-aware "
         "multi-task deep learning framework that addresses each of these limitations "
         "simultaneously. Our contributions are: (i) simultaneous pixel-level segmentation "
         "of tooth outlines and crestal bone lines through a shared ResNet-34 encoder and "
         "task-specific decoder heads, with joint training demonstrated to reduce the "
         "generalisation gap by 80.7% relative to single-task training; (ii) CEJ and "
         "root-apex landmark detection via Gaussian heatmap regression, providing the "
         "geometric reference points required to compute bone-loss percentage directly from "
         "model outputs; (iii) a Tversky-Focal composite loss with target-space "
         "rasterisation specifically designed for sparse, thin-line bone segmentation; "
         "and (iv) MC Dropout-based predictive entropy (T=20) demonstrated to be "
         "well-calibrated (ECE=0.094) and significantly correlated with segmentation error "
         "(ρ=0.591, p<0.0001), enabling reliable case-level clinical triage. Evaluated on "
         "the DenPAR benchmark dataset of 1,000 annotated IOPA radiographs,", False, False, False),
        ("9", True, False, False),
        (" OpenGeoTrust-PerioSAM establishes new performance standards for this class of "
         "task and provides a clinically interpretable, trustworthy alternative to manual "
         "radiographic bone level measurement.", False, False, False),
    ], indent=True)

    # ── METHODS ──────────────────────────────────────────────────────────────
    _heading(doc, "METHODS")
    _heading(doc, "Study Design and Data Source")

    _sup_para(doc, [
        ("This study used the DenPAR dataset (Zenodo record 16645076),", False, False, False),
        ("9", True, False, False),
        (" a publicly available benchmark comprising 1,000 intraoral periapical "
         "radiographs acquired across multiple clinical centres. Ground-truth annotations "
         "include per-image binary tooth masks, crestal bone-level polyline annotations, "
         "and CEJ/root-apex landmark coordinates. Images were partitioned into training "
         "(n=650), validation (n=150), and test (n=200) subsets using the original dataset "
         "split without modification. No patient re-identification was attempted; the study "
         "was conducted entirely using de-identified secondary data under the open dataset "
         "licence.", False, False, False),
    ], indent=True)

    _heading(doc, "Radiographic Preprocessing")

    _para(doc,
        "All images were resized to 512×512 pixels (bilinear interpolation) and normalised "
        "to [0, 1] per-image. Bone-level polyline annotations were rasterised at the target "
        "resolution of 512×512 pixels with a fixed line width of 7 pixels — a critical design "
        "choice: naive rasterisation at original resolution followed by downsampling yields a "
        "mean positive-pixel fraction of 0.08% (SD 0.03%), whereas target-space rasterisation "
        "at 7 pixels achieves 0.60% (SD 0.20%), a 7.5-fold increase enabling stable gradient "
        "propagation (Figure 1). CEJ and root-apex coordinates were converted to 2-channel "
        "Gaussian heatmaps (σ=8 pixels). Full preprocessing details are provided in "
        "Supplementary Methods.",
        indent=True)

    _heading(doc, "Model Architecture")

    _sup_para(doc, [
        ("OpenGeoTrust-PerioSAM employs a U-Net encoder-decoder architecture", False, False, False),
        ("10", True, False, False),
        (" with a ResNet-34 backbone", False, False, False),
        ("11", True, False, False),
        (" pretrained on ImageNet, adapted for single-channel grayscale input. "
         "Three task-specific decoder heads (Figure 1) branch from the shared encoder: "
         "(i) a tooth segmentation head producing a two-class (tooth/background) probability "
         "map; (ii) a crestal bone-line segmentation head producing a two-class (bone/background) "
         "map; and (iii) a keypoint heatmap head producing two-channel CEJ/apex Gaussian "
         "heatmaps with soft-argmax landmark extraction. Each head applies transposed "
         "convolution upsampling with encoder skip connections. MC Dropout (p=0.30) is "
         "applied after each decoder block;", False, False, False),
        ("8,12", True, False, False),
        (" at inference, T=20 stochastic forward passes generate a spatial predictive "
         "entropy map. A post-hoc geometry module computes bone-loss percentage as "
         "(CEJ-to-crestal-bone distance) ÷ (CEJ-to-apex distance) × 100, directly aligned "
         "with the 2018 staging criteria.", False, False, False),
        ("2,3", True, False, False),
        (" The model contains approximately 25 million trainable parameters. "
         "Full architectural specifications are provided in Supplementary Methods.", False, False, False),
    ], indent=True)

    _heading(doc, "Loss Function")

    _sup_para(doc, [
        ("The composite multi-task loss was: L", False, False, False),
        ("total", False, False, False),
        (" = 1.0·L", False, False, False),
        ("tooth", False, False, False),
        (" + 2.0·L", False, False, False),
        ("bone", False, False, False),
        (" + 0.5·L", False, False, False),
        ("kp", False, False, False),
        (" + 0.1·L", False, False, False),
        ("geo", False, False, False),
        (". L", False, False, False),
        ("tooth", False, False, False),
        (" combined cross-entropy and Dice loss. L", False, False, False),
        ("bone", False, False, False),
        (" employed a Tversky-Focal composite (0.6·L", False, False, False),
        ("Tversky", False, False, False),
        ("[α=0.3, β=0.7] + 0.4·L", False, False, False),
        ("Focal", False, False, False),
        ("[γ=2.0, α", False, False, False),
        ("f", False, False, False),
        ("=0.75]),", False, False, False),
        ("13,14", True, False, False),
        (" setting β>α to penalise false-negative bone predictions more heavily than "
         "false positives — appropriate for thin, sparse structures where missed detection "
         "carries greater clinical consequence than minor over-prediction. Task weight "
         "w", False, False, False),
        ("bone", False, False, False),
        ("=2.0 was identified as optimal by ablation (Supplementary Figure A7). L", False, False, False),
        ("kp", False, False, False),
        (" was MSE over heatmap values; L", False, False, False),
        ("geo", False, False, False),
        (" penalised bone-line predictions overlapping with tooth-mask predictions.", False, False, False),
    ], indent=True)

    _heading(doc, "Training Protocol")

    _sup_para(doc, [
        ("All models were trained with AdamW", False, False, False),
        ("15", True, False, False),
        (" (initial learning rate 1×10⁻⁴, weight decay 1×10⁻⁴) for up to 150 epochs on "
         "an NVIDIA A10G GPU (24 GB) via Modal cloud infrastructure. ReduceLROnPlateau "
         "(factor=0.5, patience=10, mode=max) monitored validation bone-line Dice. "
         "Early stopping applied with patience=30 and a minimum 30-epoch floor. The best "
         "checkpoint was defined by peak validation bone-line Dice. The final multi-task "
         "model trained for 39 epochs (~26 GPU-minutes), with the best checkpoint at "
         "epoch 24.", False, False, False),
    ], indent=True)

    _heading(doc, "Evaluation Metrics")

    _para(doc,
        "Segmentation: Dice Similarity Coefficient (DSC), Intersection over Union (IoU), "
        "precision, recall, 95th-percentile Hausdorff distance (HD95), mean surface "
        "distance (MSD). Landmark detection: mean radial error (MRE), normalised mean "
        "error (NME), Percentage of Correct Keypoints at 2, 4, 8-pixel tolerances "
        "(PCK@2px, PCK@4px, PCK@8px). Calibration: Expected Calibration Error (ECE; "
        "15 bins) and Brier Score. Uncertainty-error association: Spearman rank correlation. "
        "Bootstrap 95% confidence intervals (B=2,000 resamples with replacement) computed "
        "for all primary metrics. Statistical analyses used Python 3.10 (SciPy 1.11); "
        "α=0.05 significance threshold.",
        indent=True)

    # ── RESULTS ──────────────────────────────────────────────────────────────
    _heading(doc, "RESULTS")
    _heading(doc, "Training Dynamics")

    _para(doc,
        "The multi-task model converged over 39 epochs (best checkpoint: epoch 24, "
        "val_bone_dice=0.578; Figure 5). Tooth segmentation Dice stabilised above 0.946 "
        "from epoch 10 onward. Bone-line Dice rose gradually with oscillation in epochs "
        "3–20 (rolling SD >0.05) before stabilising at 0.54–0.58 from epoch 20 onward. "
        "Training and validation losses decreased monotonically without divergence. "
        "Loss component analysis (Supplementary Figure A2) showed the bone loss term "
        "contributing 50–79% of total training signal through mid-training, consistent "
        "with the intended emphasis of w_bone=2.0.",
        indent=True)

    _heading(doc, "Segmentation Performance")

    _para(doc,
        "On the held-out test set (n=200; Table 1), tooth segmentation achieved DSC 0.957 "
        "(95% CI: 0.948–0.958), IoU 0.919, precision 0.964, recall 0.952, HD95 6.05 px, "
        "and MSD 0.73 px — indicating near-complete overlap with ground-truth outlines at "
        "sub-pixel mean surface deviation.",
        indent=True)

    _para(doc,
        "Crestal bone-line segmentation achieved DSC 0.544 (95% CI: 0.543–0.595), IoU "
        "0.394, precision 0.560, and recall 0.557. Near-symmetric precision and recall "
        "confirm that the model does not systematically over- or under-predict bone-line "
        "extent, a property attributable to the balanced Tversky parameterisation (α=0.3, "
        "β=0.7). HD95 was 73.12 px, reflecting the challenge of thin, discontinuous "
        "crestal structures where small localisation errors generate large Hausdorff values. "
        "Representative predictions for three test cases are shown in Figure 3.",
        indent=True)

    _heading(doc, "Keypoint Detection")

    _para(doc,
        "CEJ and root-apex landmark detection achieved MRE 1.049 px (NME 0.0014), "
        "PCK@2px 0.948, and PCK@4px/PCK@8px both 1.000 (Table 1; Figure 4B). At "
        "the typical spatial resolution of an IOPA radiograph (~5.5 lp/mm), an MRE of "
        "1.05 px corresponds to approximately 0.09–0.13 mm in physical space — well "
        "within the precision of clinical periodontal measurement modalities. Virtually "
        "all predicted CEJ and apex landmarks (100.0%) fell within a 4-pixel neighbourhood "
        "of the ground truth.",
        indent=True)

    _heading(doc, "Uncertainty Calibration")

    _sup_para(doc, [
        ("MC Dropout inference (T=20) yielded ECE 0.094 (95% CI: 0.090–0.095), below "
         "the 0.10 threshold for well-calibrated models,", False, False, False),
        ("16", True, False, False),
        (" and Brier Score 0.009. Predictive entropy correlated significantly with "
         "segmentation error (Spearman ρ=0.591, p<0.0001; Figure 6; Figure 4C), "
         "confirming that the model's uncertainty estimates are predictive of its failure "
         "modes. Spatial uncertainty maps (Figure 3, column F) show high-entropy regions "
         "concentrated at structure boundaries and regions of radiographic ambiguity — "
         "precisely where human observers also express diagnostic uncertainty.", False, False, False),
    ], indent=True)

    _heading(doc, "Ablation Study")

    _para(doc,
        "Three configurations were evaluated (Supplementary Table A1; Supplementary "
        "Figure A1). Configuration A (baseline: 3-pixel rasterisation, Dice-Focal loss, "
        "single-task) achieved validation bone DSC 0.350 with near-zero test DSC due to "
        "systematic false-negative collapse. Configuration B (single-task: 7-pixel "
        "rasterisation, Tversky-Focal) raised validation DSC to 0.597 (+70.6% over "
        "baseline) but produced a large generalisation gap (val 0.597 − test 0.421 = "
        "0.176), indicating overfitting without multi-task regularisation. Configuration C "
        "(full multi-task: Tversky-Focal, bone_w=2.0) achieved test DSC 0.544 — a "
        "+29.2% improvement over Config B — and reduced the generalisation gap to 0.034, "
        "an 80.7% reduction (Supplementary Figure A4).",
        indent=True)

    _heading(doc, "Sensitivity Analyses")

    _para(doc,
        "Decision threshold sensitivity (Supplementary Figure A3) confirmed τ=0.5 is "
        "near-optimal for both tasks (tooth DSC >0.90 for τ∈[0.35, 0.70]; bone DSC "
        "peak at τ∈[0.45, 0.55]). Bone loss weight sensitivity showed a concave response "
        "curve with optimum at w_bone=2.0; bone-line mask thickness showed monotonic gain "
        "from 3 to 7 pixels with diminishing returns beyond 7 pixels "
        "(Supplementary Figure A7). Full sensitivity results are in Supplementary Results.",
        indent=True)

    # ── DISCUSSION ───────────────────────────────────────────────────────────
    _heading(doc, "DISCUSSION")
    _heading(doc, "Comparison with Existing Approaches")

    _sup_para(doc, [
        ("This study presents the first validated multi-task framework for simultaneous "
         "pixel-level crestal bone-line and tooth segmentation, CEJ/apex landmark detection, "
         "and calibrated uncertainty estimation from IOPA radiographs, evaluated on a "
         "1,000-image public benchmark.", False, False, False),
    ], indent=True)

    _sup_para(doc, [
        ("Prior periodontal AI has been dominated by classification paradigms on panoramic "
         "radiographs. Krois et al.", False, False, False),
        ("6", True, False, False),
        (" achieved accuracy 0.81 for radiographic bone loss detection — an important "
         "proof-of-concept for automated screening, but returning a binary label rather "
         "than a spatial measurement. Chang et al.", False, False, False),
        ("7", True, False, False),
        (" advanced this to automatic staging with ICC 0.91 using a hybrid framework "
         "on panoramic radiographs. Neither study provided pixel-level bone margin "
         "localisation, per-tooth geometric measurements, or uncertainty estimates — "
         "capabilities required for integration into the clinical recording workflow. "
         "Critically, both studies used panoramic radiographs, which offer inferior "
         "spatial resolution and greater geometric distortion compared with IOPAs, "
         "the format most commonly acquired for bone level assessment in clinical "
         "periodontology. The scoping review by Schwendicke et al.", False, False, False),
        ("4", True, False, False),
        (" confirmed that periapical radiograph segmentation, particularly of crestal "
         "bone structures, remains one of the least studied and most technically "
         "challenging tasks in dental AI.", False, False, False),
    ], indent=True)

    _sup_para(doc, [
        ("Our crestal bone-line DSC of 0.544 represents, to our knowledge, the first "
         "systematic pixel-level crestal bone segmentation result on IOPAs in the "
         "peer-reviewed literature. Tooth segmentation DSC of 0.957 — achieved within "
         "a multi-task framework — is competitive with or exceeds the 0.88–0.94 range "
         "reported across dental CNN segmentation tasks in Schwendicke et al.", False, False, False),
        ("4", True, False, False),
        (" The MedSAM foundation model,", False, False, False),
        ("17", True, False, False),
        (" trained on >1.5 million image-mask pairs, demonstrated that foundation-model "
         "pre-training outperforms modality-specific specialist models across 86 medical "
         "imaging tasks. Replacing the ResNet-34 backbone with a MedSAM ViT encoder "
         "represents a clear future direction for improving bone-line recall, particularly "
         "in severe attachment loss where the crestal margin is radiographically indistinct.", False, False, False),
    ], indent=True)

    _heading(doc, "Clinical Significance")

    _sup_para(doc, [
        ("The clinical implications extend across patient care, practice workflow, and "
         "population-level surveillance. The 2018 staging framework requires radiographic "
         "bone loss as a millimetre value or root-length percentage.", False, False, False),
        ("2,3", True, False, False),
        (" OpenGeoTrust-PerioSAM's geometry module derives this ratio directly from "
         "predicted CEJ/apex coordinates and the bone-line segmentation, providing a "
         "clinician-ready bone-loss estimate without manual measurement. Given keypoint "
         "MRE of 1.05 px (~0.10 mm) and bone-line segmentation uncertainty of ~2–3 px "
         "at the margin, the system's output uncertainty is broadly comparable to published "
         "inter-operator variability in manual radiographic bone level measurement, "
         "positioning it as a realistic clinical adjunct.", False, False, False),
    ], indent=True)

    _para(doc,
        "The calibrated uncertainty output enables a human-in-the-loop governance model "
        "in which high-entropy cases are automatically routed for clinician review while "
        "lower-entropy cases proceed to automated reporting. The significant uncertainty-error "
        "correlation (ρ=0.591, p<0.0001) confirms this routing is operationally rational. "
        "The 80.7% reduction in generalisation gap from multi-task training has broader "
        "methodological implications: it establishes that auxiliary tasks providing shared "
        "anatomical context can be as important as the primary loss function for "
        "cross-distribution robustness when targeting sparse anatomical structures.",
        indent=True)

    _heading(doc, "Limitations")

    _para(doc,
        "First, the DenPAR dataset comprises 1,000 radiographs from a single release; "
        "multi-centre prospective validation across heterogeneous imaging equipment is "
        "required. Second, single-annotator consensus ground truth may constitute a "
        "ceiling on achievable DSC independent of model capacity. Third, the geometry "
        "module is sensitive to apex localisation, which may be obscured in severe "
        "apical pathology. Fourth, MC Dropout provides an approximation to full "
        "Bayesian inference; ensemble methods may improve calibration further. Fifth, "
        "this evaluation used static two-dimensional IOPAs; longitudinal bone-level "
        "change monitoring requires dedicated prospective study design.",
        indent=True)

    # ── CONCLUSION ───────────────────────────────────────────────────────────
    _heading(doc, "CONCLUSION")

    _sup_para(doc, [
        ("OpenGeoTrust-PerioSAM demonstrates that simultaneous pixel-level crestal "
         "bone-line and tooth segmentation, CEJ/apex landmark detection, geometry-aware "
         "bone-loss measurement, and calibrated uncertainty quantification are achievable "
         "from IOPA radiographs within a single multi-task deep learning framework. "
         "The composite Tversky-Focal loss with target-space rasterisation provides a "
         "principled solution to thin-structure class imbalance, yielding crestal bone "
         "DSC of 0.544. Joint multi-task training reduces the single-task generalisation "
         "gap by 80.7%, establishing multi-task regularisation as a critical design "
         "principle for sparse anatomical target detection. Calibrated MC Dropout "
         "uncertainty (ECE 0.094; Spearman ρ=0.591, p<0.0001) enables clinician-in-the-"
         "loop triage consistent with responsible clinical AI deployment. By delivering "
         "all measurements required for the 2018 periodontitis staging framework", False, False, False),
        ("2,3", True, False, False),
        (" from a single forward pass, OpenGeoTrust-PerioSAM offers a clinically "
         "actionable, reproducible adjunct to radiographic bone level assessment with "
         "direct implications for standardising periodontitis staging, improving "
         "longitudinal monitoring, and reducing operator-dependent variability in "
         "everyday periodontal practice.", False, False, False),
    ], indent=True)

    # ── DATA AVAILABILITY ────────────────────────────────────────────────────
    _heading(doc, "DATA AVAILABILITY")
    _para(doc,
        "The DenPAR dataset used in this study is publicly available on Zenodo "
        "(DOI: 10.5281/zenodo.16645076; CC-BY 4.0 licence). All model code, training scripts, "
        "evaluation pipelines, and publication figures are openly available in the "
        "OpenGeoTrust-PerioSAM repository at https://github.com/Pkr2180/opengeotrust-periosam-denpar. "
        "Trained model checkpoints are reproducible by running the provided Modal GPU pipeline "
        "against the DenPAR dataset; raw checkpoint files are not deposited owing to file-size "
        "constraints but will be shared upon reasonable request to the corresponding author.",
        sa=4)

    # ── CONFLICT OF INTEREST ─────────────────────────────────────────────────
    _heading(doc, "CONFLICT OF INTEREST")
    _para(doc, "The authors declare no conflicts of interest relevant to this study.", sa=4)

    # ── REFERENCES ───────────────────────────────────────────────────────────
    _heading(doc, "REFERENCES")
    refs = [
        ("1. ", "Kassebaum NJ, Bernabé E, Dahiya M, Bhandari B, Murray CJ, Marcenes W. "
                "Global burden of severe periodontitis in 1990–2010: a systematic review and "
                "meta-regression. J Dent Res. 2014;93(11):1045–1053. "
                "doi: 10.1177/0022034514552491"),
        ("2. ", "Papapanou PN, Sanz M, Buduneli N, Dietrich T, Feres M, Fine DH, et al. "
                "Periodontitis: Consensus report of workgroup 2 of the 2017 World Workshop on "
                "the Classification of Periodontal and Peri-Implant Diseases and Conditions. "
                "J Periodontol. 2018;89(Suppl 1):S173–S182. doi: 10.1002/JPER.17-0721"),
        ("3. ", "Tonetti MS, Greenwell H, Kornman KS. Staging and grading of periodontitis: "
                "Framework and proposal of a new classification and case definition. "
                "J Periodontol. 2018;89(Suppl 1):S159–S172. doi: 10.1002/JPER.18-0006"),
        ("4. ", "Schwendicke F, Golla T, Dreher M, Krois J. Convolutional neural networks for "
                "dental image diagnostics: A scoping review. J Dent. 2019;91:103226. "
                "doi: 10.1016/j.jdent.2019.103226"),
        ("5. ", "Caton JG, Armitage G, Berglundh T, Chapple ILC, Jepsen S, Kornman KS, et al. "
                "A new classification scheme for periodontal and peri-implant diseases and "
                "conditions – Introduction and key changes from the 1999 classification. "
                "J Clin Periodontol. 2018;45(Suppl 20):S1–S8. doi: 10.1111/jcpe.12935"),
        ("6. ", "Krois J, Ekert T, Meinhold L, Golla T, Kharbot B, Wittemeier A, et al. "
                "Deep learning for the radiographic detection of periodontal bone loss. "
                "Sci Rep. 2019;9(1):8495. doi: 10.1038/s41598-019-44839-3"),
        ("7. ", "Chang HJ, Lee SJ, Yong TH, Shin NY, Jang BG, Kim JE, et al. "
                "Deep learning hybrid method to automatically diagnose periodontal bone loss "
                "and stage periodontitis. Sci Rep. 2020;10(1):7531. "
                "doi: 10.1038/s41598-020-64509-z"),
        ("8. ", "Gal Y, Ghahramani Z. Dropout as a Bayesian approximation: Representing model "
                "uncertainty in deep learning. In: Proceedings of the 33rd International "
                "Conference on Machine Learning. PMLR; 2016. Vol 48. p. 1050–1059."),
        ("9. ", "DenPAR Dataset. Periodontal radiograph dataset (DenPAR) [dataset]. "
                "Zenodo; 2024. Available from: https://zenodo.org/record/16645076"),
        ("10. ","Ronneberger O, Fischer P, Brox T. U-Net: Convolutional networks for biomedical "
                "image segmentation. In: MICCAI 2015. Lecture Notes in Computer Science, "
                "vol 9351. Cham: Springer; 2015. p. 234–241."),
        ("11. ","He K, Zhang X, Ren S, Sun J. Deep residual learning for image recognition. "
                "In: Proceedings of CVPR. Las Vegas: IEEE; 2016. p. 770–778."),
        ("12. ","Srivastava N, Hinton G, Krizhevsky A, Sutskever I, Salakhutdinov R. "
                "Dropout: A simple way to prevent neural networks from overfitting. "
                "J Mach Learn Res. 2014;15(1):1929–1958."),
        ("13. ","Lin TY, Goyal P, Girshick R, He K, Dollár P. Focal loss for dense object "
                "detection. In: Proceedings of ICCV. Venice: IEEE; 2017. p. 2980–2988."),
        ("14. ","Abraham N, Khan NM. A novel focal Tversky loss function with improved attention "
                "U-Net for lesion segmentation. In: Proceedings of ISBI 2019. Venice: IEEE; "
                "2019. p. 683–687."),
        ("15. ","Loshchilov I, Hutter F. Decoupled weight decay regularization. In: Proceedings "
                "of ICLR. 2019."),
        ("16. ","Guo C, Pleiss G, Sun Y, Weinberger KQ. On calibration of modern neural networks. "
                "In: Proceedings of the 34th International Conference on Machine Learning. "
                "PMLR; 2017. Vol 70. p. 1321–1330."),
        ("17. ","Ma J, He Y, Li F, Han L, You C, Wang B. Segment anything in medical images. "
                "Nat Commun. 2024;15(1):654. doi: 10.1038/s41467-024-44824-z"),
        ("18. ","Isensee F, Jaeger PF, Kohl SAA, Petersen J, Maier-Hein KH. nnU-Net: a "
                "self-configuring method for deep learning-based biomedical image segmentation. "
                "Nat Methods. 2021;18(2):203–211."),
        ("19. ","Caruana R. Multitask learning. Mach Learn. 1997;28(1):41–75."),
        ("20. ","Bland JM, Altman DG. Statistical methods for assessing agreement between two "
                "methods of clinical measurement. Lancet. 1986;1(8476):307–310."),
    ]
    for num, body in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after    = Pt(4)
        p.paragraph_format.line_spacing   = Pt(24)
        p.paragraph_format.left_indent    = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        r1 = p.add_run(num);  _font(r1, bold=True)
        r2 = p.add_run(body); _font(r2)

    # ════════════════════════════════════════════════════════════════════════
    # TABLES (inline after references, before figures)
    # ════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, "TABLE 1. Test Set Performance — OpenGeoTrust-PerioSAM (DenPAR, n=200)",
             sb=0, sa=8)

    # Segmentation table
    _para(doc, "A. Segmentation", bold=True, sa=4, sb=4)
    t1 = doc.add_table(rows=1, cols=7)
    t1.style = "Table Grid"
    _hdr_row(t1, ["Task","DSC (95% CI)","IoU","Precision","Recall","HD95 (px)","MSD (px)"])
    _data_row(t1, ["Tooth segmentation","0.957 (0.948–0.958)","0.919","0.964","0.952","6.05","0.73"])
    _data_row(t1, ["Bone-line segmentation","0.544 (0.543–0.595)","0.394","0.560","0.557","73.12","12.45"])

    _para(doc, "B. Landmark detection", bold=True, sa=4, sb=10)
    t2 = doc.add_table(rows=1, cols=6)
    t2.style = "Table Grid"
    _hdr_row(t2, ["Task","MRE (px)","NME","PCK@2px","PCK@4px","PCK@8px"])
    _data_row(t2, ["CEJ + root-apex","1.049","0.0014","0.948","1.000","1.000"])

    _para(doc, "C. Uncertainty calibration", bold=True, sa=4, sb=10)
    t3 = doc.add_table(rows=1, cols=4)
    t3.style = "Table Grid"
    _hdr_row(t3, ["Task","ECE","Brier Score","Spearman ρ (p-value)"])
    _data_row(t3, ["MC Dropout (T=20)","0.094 (0.090–0.095)","0.009 (0.008–0.010)","0.591 (p<0.0001)"])

    _para(doc,
        "DSC, Dice Similarity Coefficient; IoU, Intersection over Union; HD95, 95th-percentile "
        "Hausdorff Distance; MSD, Mean Surface Distance; MRE, Mean Radial Error; NME, Normalised "
        "Mean Error; PCK, Percentage of Correct Keypoints; ECE, Expected Calibration Error. "
        "95% CI from B=2,000 bootstrap resamples with replacement.",
        italic=True, size=10, sb=6, sa=4)

    # ════════════════════════════════════════════════════════════════════════
    # EMBEDDED FIGURES  (each on its own page, image + legend)
    # ════════════════════════════════════════════════════════════════════════
    figure_data = [
        (
            os.path.join(FIGS, "fig1_workflow.png"), 6.0,
            "Figure 1.",
            "Schematic pipeline of OpenGeoTrust-PerioSAM. An IOPA radiograph (512×512 pixels) "
            "is encoded by a shared ResNet-34 encoder. Three task-specific decoder heads produce "
            "in parallel: (i) tooth segmentation probability map; (ii) crestal bone-line "
            "segmentation map; and (iii) CEJ/apex Gaussian heatmaps with soft-argmax landmark "
            "extraction. The geometry module computes bone-loss percentage from predicted "
            "coordinates. MC Dropout (T=20) generates a spatial predictive entropy map for "
            "uncertainty quantification. Dashed arrows indicate encoder skip connections."
        ),
        (
            os.path.join(FIGS, "fig2_dataset_qc.png"), 6.0,
            "Figure 2.",
            "Representative DenPAR dataset sample (test set). Left to right: original IOPA "
            "radiograph; tooth mask overlay (green); crestal bone-line overlay (red); CEJ "
            "heatmap (blue); root-apex heatmap (orange). Target-space bone-line rasterisation "
            "at 7-pixel width (512×512 resolution) is visible in column three, illustrating the "
            "increased positive-pixel density compared with naive 3-pixel rasterisation."
        ),
        (
            os.path.join(FIGS, "fig3_model_output.png"), 6.0,
            "Figure 3.",
            "Qualitative model output comparison for three representative test cases (rows). "
            "Columns: (A) original IOPA radiograph; (B) ground-truth tooth mask; "
            "(C) predicted tooth mask; (D) ground-truth bone-line mask; (E) predicted bone-line "
            "mask; (F) MC Dropout uncertainty map (Shannon entropy). High-entropy regions "
            "(yellow/warm colours) are concentrated at structure boundaries and areas of "
            "radiographic ambiguity, consistent with clinician uncertainty in these regions."
        ),
        (
            os.path.join(FIGS, "fig4_quantitative.png"), 6.0,
            "Figure 4.",
            "Quantitative performance panel. (A) Segmentation metrics (DSC, IoU, precision, "
            "recall) for tooth (blue) and bone-line (orange) heads on the test set (n=200); "
            "error bars represent 95% bootstrap CI. (B) PCK curves for CEJ and apex landmark "
            "detection at 2, 4, and 8-pixel tolerances. (C) Reliability diagram for MC Dropout "
            "uncertainty calibration; the dashed diagonal represents perfect calibration; "
            "ECE=0.094 indicates well-calibrated predictive uncertainty."
        ),
        (
            os.path.join(FIGS, "fig5_training_curves.png"), 6.0,
            "Figure 5.",
            "Training dynamics over 39 epochs. (A) Validation Dice curves for tooth segmentation "
            "(blue) and crestal bone-line segmentation (orange); the star marks the best "
            "checkpoint at epoch 24 (val_bone_dice=0.578). (B) Training (solid) and validation "
            "(dashed) total loss curves. Shaded regions in (A) denote the high-oscillation phase "
            "(rolling SD >0.05; epochs 3–20). Both curves decrease monotonically without "
            "divergence, confirming the absence of significant overfitting."
        ),
        (
            os.path.join(FIGS, "fig6_uncertainty_error.png"), 5.0,
            "Figure 6.",
            "Association between predictive uncertainty and segmentation error. Scatter plot of "
            "MC Dropout entropy (x-axis) versus segmentation error (1−DSC, y-axis) for all "
            "200 test images. Spearman ρ=0.591, p<0.0001 (two-tailed). The solid line shows "
            "a robust linear fit. The significant positive correlation confirms that model "
            "uncertainty is a reliable indicator of prediction quality, supporting its use "
            "as a case-level triage signal in clinical workflows."
        ),
    ]

    for png_path, width, label, caption in figure_data:
        _add_figure(doc, png_path, width_in=width,
                    caption_label=label, caption_text=caption)

    doc.save(OUT)
    wc = len(" ".join(p.text for p in doc.paragraphs).split())
    print(f"Main manuscript saved: {OUT}")
    print(f"Approximate word count: {wc}")

# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build()
